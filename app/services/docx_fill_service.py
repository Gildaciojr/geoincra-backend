# app/services/docx_fill_service.py
from __future__ import annotations

from pathlib import Path
from typing import Any
import re
import shutil
from datetime import datetime

from docx import Document  # python-docx


PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_.-]+)\s*}}")


def _get_value(data: dict[str, Any], key: str) -> str:
    """
    Suporta chaves simples e com ponto:
      - "cliente_nome"
      - "proprietario.cpf"
    """
    if "." not in key:
        val = data.get(key, "")
        return "" if val is None else str(val)

    cur: Any = data
    for part in key.split("."):
        if not isinstance(cur, dict):
            return ""
        cur = cur.get(part)
        if cur is None:
            return ""
    return str(cur)


def _replace_in_paragraph(paragraph, data: dict[str, Any]) -> None:
    """
    Estratégia: junta texto de runs, faz replace, reescreve em um run.
    Isso resolve a maioria dos casos de run-splitting.
    """
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    def repl(match):
        key = match.group(1)
        return _get_value(data, key)

    new_text = PLACEHOLDER_RE.sub(repl, full_text)

    # limpa runs e coloca tudo no primeiro
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def _replace_in_table(table, data: dict[str, Any]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, data)
            for nested in cell.tables:
                _replace_in_table(nested, data)


def fill_docx_template(template_path: Path, data: dict[str, Any], output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)

    if not template_path.exists():
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    # cria cópia para escrita segura
    ts = int(datetime.utcnow().timestamp())
    out_path = output_dir / f"requerimento_{ts}.docx"
    shutil.copyfile(template_path, out_path)

    doc = Document(str(out_path))

    # paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p, data)

    # tables
    for t in doc.tables:
        _replace_in_table(t, data)

    doc.save(str(out_path))
    return out_path
